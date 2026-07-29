from __future__ import annotations

from io import BytesIO
from pathlib import Path
from typing import Any
from zipfile import BadZipFile, ZipFile

import pypdfium2 as pdfium
from docx import Document
from openpyxl import load_workbook
from openpyxl.utils import get_column_letter
from pypdf import PdfReader

from coal_platform.ocr import OCRBackend

MAX_FILE_BYTES = 50 * 1024 * 1024
MAX_EXPANDED_BYTES = 200 * 1024 * 1024
MAX_ARCHIVE_ENTRIES = 10000
MAX_PAGES = 1000
MAX_BLOCKS = 10000
MAX_BLOCK_CHARS = 100000
MAX_OCR_PAGE_PIXELS = 40_000_000
MAX_REGION_PIXELS = 40_000_000
THUMBNAIL_MAX_PIXELS = 2_000_000


class DocumentParseError(ValueError):
    pass


def _block(
    page_no: int,
    block_type: str,
    content_text: str,
    source_ref: str,
    *,
    confidence: float = 1.0,
    bbox: dict[str, Any] | None = None,
) -> dict[str, Any]:
    return {
        "page_no": page_no,
        "block_type": block_type,
        "content_text": content_text[:MAX_BLOCK_CHARS],
        "bbox": bbox,
        "confidence": confidence,
        "source_ref": source_ref,
    }


def _ensure_limits(content: bytes, blocks: list[dict[str, Any]]) -> None:
    if len(content) > MAX_FILE_BYTES:
        raise DocumentParseError(f"file exceeds {MAX_FILE_BYTES} byte parse limit")
    if len(blocks) > MAX_BLOCKS:
        raise DocumentParseError(f"document exceeds {MAX_BLOCKS} block parse limit")


def _ensure_archive_limits(content: bytes) -> None:
    try:
        with ZipFile(BytesIO(content)) as archive:
            entries = archive.infolist()
            if len(entries) > MAX_ARCHIVE_ENTRIES:
                raise DocumentParseError(f"document archive exceeds {MAX_ARCHIVE_ENTRIES} entry limit")
            expanded_size = sum(item.file_size for item in entries)
            if expanded_size > MAX_EXPANDED_BYTES:
                raise DocumentParseError(f"document archive exceeds {MAX_EXPANDED_BYTES} expanded byte limit")
    except BadZipFile as exc:
        raise DocumentParseError("invalid OOXML document archive") from exc


def _pdf_page_has_images(page: Any) -> bool:
    try:
        resources = page.get("/Resources")
        resources = resources.get_object() if hasattr(resources, "get_object") else resources
        xobjects = resources.get("/XObject") if resources else None
        xobjects = xobjects.get_object() if hasattr(xobjects, "get_object") else xobjects
        return any(
            (item.get_object() if hasattr(item, "get_object") else item).get("/Subtype") == "/Image"
            for item in (xobjects or {}).values()
        )
    except (AttributeError, TypeError, ValueError):
        return False


def _ocr_pdf_pages(
    content: bytes,
    page_numbers: list[int],
    ocr_backend: OCRBackend,
    dpi: int,
) -> tuple[list[dict[str, Any]], list[int]]:
    blocks: list[dict[str, Any]] = []
    unresolved_pages: list[int] = []
    try:
        document = pdfium.PdfDocument(content)
    except Exception as exc:
        raise DocumentParseError(f"failed to render PDF for OCR: {exc}") from exc
    try:
        for page_no in page_numbers:
            page = document.get_page(page_no - 1)
            bitmap = None
            image = None
            try:
                page_width, page_height = page.get_size()
                estimated_pixels = page_width * dpi / 72 * page_height * dpi / 72
                if estimated_pixels > MAX_OCR_PAGE_PIXELS:
                    raise DocumentParseError(f"PDF page {page_no} exceeds {MAX_OCR_PAGE_PIXELS} OCR pixel limit")
                bitmap = page.render(scale=dpi / 72)
                image = bitmap.to_pil()
                image_output = BytesIO()
                image.save(image_output, format="PNG")
                lines = ocr_backend.recognize(image_output.getvalue())
                if not lines:
                    unresolved_pages.append(page_no)
                    continue
                scale_x = page_width / image.width
                scale_y = page_height / image.height
                for line_no, line in enumerate(lines, start=1):
                    bbox = {
                        "x": round(line.left * scale_x, 3),
                        "y": round(line.top * scale_y, 3),
                        "width": round(line.width * scale_x, 3),
                        "height": round(line.height * scale_y, 3),
                        "page_width": round(page_width, 3),
                        "page_height": round(page_height, 3),
                        "unit": "pt",
                    }
                    blocks.append(
                        _block(
                            page_no,
                            "ocr_line",
                            line.text,
                            f"ocr:{ocr_backend.engine_name}:page:{page_no}:line:{line_no}",
                            confidence=line.confidence,
                            bbox=bbox,
                        )
                    )
                    if len(blocks) > MAX_BLOCKS:
                        raise DocumentParseError(f"OCR result exceeds {MAX_BLOCKS} block parse limit")
            finally:
                if image:
                    image.close()
                if bitmap:
                    bitmap.close()
                page.close()
    except Exception as exc:
        raise DocumentParseError(f"PDF OCR failed: {exc}") from exc
    finally:
        document.close()
    return blocks, unresolved_pages


def _parse_pdf(
    content: bytes,
    file_name: str,
    file_type: str | None = None,
    ocr_backend: OCRBackend | None = None,
    ocr_dpi: int = 200,
) -> tuple[list[dict[str, Any]], dict[str, Any]]:
    try:
        reader = PdfReader(BytesIO(content))
    except Exception as exc:
        raise DocumentParseError(f"invalid PDF document: {exc}") from exc
    if reader.is_encrypted:
        try:
            if reader.decrypt("") == 0:
                raise DocumentParseError("encrypted PDF requires a password")
        except DocumentParseError:
            raise
        except Exception as exc:
            raise DocumentParseError("encrypted PDF cannot be opened") from exc
    if len(reader.pages) > MAX_PAGES:
        raise DocumentParseError(f"PDF exceeds {MAX_PAGES} page parse limit")
    blocks: list[dict[str, Any]] = []
    empty_pages: list[int] = []
    drawing_pages: list[int] = []
    drawing_page_details: list[dict[str, Any]] = []
    drawing_identity = f"{file_name} {file_type or ''}".lower()
    drawing_name = any(keyword in drawing_identity for keyword in ("图纸", "drawing", "cad", "dwg"))
    for page_no, page in enumerate(reader.pages, start=1):
        try:
            text = page.extract_text() or ""
        except Exception as exc:
            raise DocumentParseError(f"failed to extract PDF page {page_no}: {exc}") from exc
        paragraphs = [item.strip() for item in text.split("\n") if item.strip()]
        keyword_page = any(keyword in text.lower() for keyword in ("图纸", "标题栏", "drawing", "title block", "dwg", "cad"))
        page_width = float(page.mediabox.width)
        page_height = float(page.mediabox.height)
        landscape_image_page = _pdf_page_has_images(page) and page_width > page_height * 1.2
        if drawing_name:
            detection_method, detection_confidence = "drawing_filename", 0.95
        elif keyword_page:
            detection_method, detection_confidence = "drawing_text_keyword", 0.9
        elif landscape_image_page:
            detection_method, detection_confidence = "landscape_image_page", 0.65
        else:
            detection_method, detection_confidence = "not_detected", 0.0
        if detection_confidence:
            drawing_pages.append(page_no)
            drawing_page_details.append(
                {
                    "page_no": page_no,
                    "detection_method": detection_method,
                    "detection_confidence": detection_confidence,
                }
            )
        if not paragraphs:
            empty_pages.append(page_no)
        for index, paragraph in enumerate(paragraphs, start=1):
            blocks.append(_block(page_no, "paragraph", paragraph, f"pdf:page:{page_no}:line:{index}"))
    ocr_blocks: list[dict[str, Any]] = []
    unresolved_pages = empty_pages
    if empty_pages and ocr_backend:
        ocr_blocks, unresolved_pages = _ocr_pdf_pages(content, empty_pages, ocr_backend, ocr_dpi)
        blocks.extend(ocr_blocks)
    return blocks, {
        "parser": "pypdf",
        "page_count": len(reader.pages),
        "empty_text_pages": empty_pages,
        "ocr_engine": ocr_backend.engine_name if ocr_backend and empty_pages else None,
        "ocr_page_count": len(empty_pages) - len(unresolved_pages),
        "ocr_block_count": len(ocr_blocks),
        "unresolved_ocr_pages": unresolved_pages,
        "needs_ocr": bool(unresolved_pages),
        "drawing_pages": drawing_pages,
        "drawing_page_count": len(drawing_pages),
        "drawing_page_details": drawing_page_details,
    }


def render_pdf_page_assets(
    content: bytes,
    drawing_page_details: list[dict[str, Any]],
    dpi: int = 120,
) -> list[dict[str, Any]]:
    try:
        document = pdfium.PdfDocument(content)
    except Exception as exc:
        raise DocumentParseError(f"failed to render PDF thumbnails: {exc}") from exc
    assets: list[dict[str, Any]] = []
    drawing_by_page = {item["page_no"]: item for item in drawing_page_details}
    try:
        for page_no in range(1, len(document) + 1):
            page = document.get_page(page_no - 1)
            bitmap = None
            image = None
            try:
                page_width, page_height = page.get_size()
                scale = dpi / 72
                estimated_width = max(1, round(page_width * scale))
                estimated_height = max(1, round(page_height * scale))
                if estimated_width * estimated_height > THUMBNAIL_MAX_PIXELS:
                    scale *= (THUMBNAIL_MAX_PIXELS / (estimated_width * estimated_height)) ** 0.5
                bitmap = page.render(scale=scale)
                image = bitmap.to_pil()
                output = BytesIO()
                image.save(output, format="PNG", optimize=True)
                drawing = drawing_by_page.get(page_no)
                assets.append(
                    {
                        "page_no": page_no,
                        "content": output.getvalue(),
                        "page_width": round(page_width, 3),
                        "page_height": round(page_height, 3),
                        "is_drawing": drawing is not None,
                        "detection_method": drawing["detection_method"] if drawing else "not_detected",
                        "detection_confidence": drawing["detection_confidence"] if drawing else 0.0,
                    }
                )
            finally:
                if image:
                    image.close()
                if bitmap:
                    bitmap.close()
                page.close()
    finally:
        document.close()
    return assets


def render_pdf_region(
    content: bytes,
    page_no: int,
    bbox: dict[str, float],
    dpi: int = 200,
) -> bytes:
    try:
        document = pdfium.PdfDocument(content)
        page = document.get_page(page_no - 1)
    except Exception as exc:
        raise DocumentParseError(f"failed to open PDF region: {exc}") from exc
    bitmap = None
    image = None
    try:
        page_width, page_height = page.get_size()
        x = max(0.0, min(float(bbox.get("x", 0)), page_width))
        y = max(0.0, min(float(bbox.get("y", 0)), page_height))
        width = max(0.0, min(float(bbox.get("width", 0)), page_width - x))
        height = max(0.0, min(float(bbox.get("height", 0)), page_height - y))
        if width <= 0 or height <= 0:
            raise DocumentParseError("region bbox must have positive dimensions")
        scale = dpi / 72
        if width * scale * height * scale > MAX_REGION_PIXELS:
            raise DocumentParseError(f"region exceeds {MAX_REGION_PIXELS} pixel limit")
        bitmap = page.render(scale=scale)
        image = bitmap.to_pil()
        crop = image.crop(
            (
                round(x * image.width / page_width),
                round(y * image.height / page_height),
                round((x + width) * image.width / page_width),
                round((y + height) * image.height / page_height),
            )
        )
        output = BytesIO()
        crop.save(output, format="PNG", optimize=True)
        crop.close()
        return output.getvalue()
    except DocumentParseError:
        raise
    except Exception as exc:
        raise DocumentParseError(f"failed to crop PDF region: {exc}") from exc
    finally:
        if image:
            image.close()
        if bitmap:
            bitmap.close()
        page.close()
        document.close()


def _parse_docx(content: bytes) -> tuple[list[dict[str, Any]], dict[str, Any]]:
    _ensure_archive_limits(content)
    try:
        document = Document(BytesIO(content))
    except Exception as exc:
        raise DocumentParseError(f"invalid Word document: {exc}") from exc
    blocks: list[dict[str, Any]] = []
    for index, paragraph in enumerate(document.paragraphs, start=1):
        text = paragraph.text.strip()
        if text:
            blocks.append(_block(1, "paragraph", text, f"docx:paragraph:{index}"))
    table_rows = 0
    for table_no, table in enumerate(document.tables, start=1):
        for row_no, row in enumerate(table.rows, start=1):
            values = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            if not any(values):
                continue
            table_rows += 1
            blocks.append(_block(1, "table_row", "\t".join(values), f"docx:table:{table_no}:row:{row_no}"))
    return blocks, {
        "parser": "python-docx",
        "logical_page_count": 1,
        "paragraph_count": sum(item["block_type"] == "paragraph" for item in blocks),
        "table_count": len(document.tables),
        "table_row_count": table_rows,
        "needs_ocr": False,
    }


def _parse_xlsx(content: bytes) -> tuple[list[dict[str, Any]], dict[str, Any]]:
    _ensure_archive_limits(content)
    try:
        workbook = load_workbook(BytesIO(content), read_only=True, data_only=True)
    except Exception as exc:
        raise DocumentParseError(f"invalid Excel workbook: {exc}") from exc
    blocks: list[dict[str, Any]] = []
    try:
        if len(workbook.worksheets) > MAX_PAGES:
            raise DocumentParseError(f"workbook exceeds {MAX_PAGES} sheet parse limit")
        for sheet_no, worksheet in enumerate(workbook.worksheets, start=1):
            for row_no, row in enumerate(worksheet.iter_rows(), start=1):
                values = ["" if cell.value is None else str(cell.value) for cell in row]
                while values and not values[-1]:
                    values.pop()
                if not any(values):
                    continue
                end_column = get_column_letter(max(1, len(values)))
                source_ref = f"xlsx:{worksheet.title}!A{row_no}:{end_column}{row_no}"
                blocks.append(_block(sheet_no, "table_row", "\t".join(values), source_ref))
                if len(blocks) > MAX_BLOCKS:
                    raise DocumentParseError(f"workbook exceeds {MAX_BLOCKS} row parse limit")
    finally:
        workbook.close()
    return blocks, {
        "parser": "openpyxl",
        "sheet_count": len(workbook.sheetnames),
        "sheet_names": workbook.sheetnames,
        "needs_ocr": False,
    }


def _parse_text(content: bytes) -> tuple[list[dict[str, Any]], dict[str, Any]]:
    try:
        text = content.decode("utf-8-sig")
    except UnicodeDecodeError:
        try:
            text = content.decode("gb18030")
        except UnicodeDecodeError as exc:
            raise DocumentParseError("text document encoding must be UTF-8 or GB18030") from exc
    blocks = [
        _block(1, "paragraph", line.strip(), f"text:line:{line_no}")
        for line_no, line in enumerate(text.splitlines(), start=1)
        if line.strip()
    ]
    return blocks, {"parser": "plain-text", "logical_page_count": 1, "needs_ocr": False}


def parse_document(
    content: bytes,
    file_name: str,
    file_type: str | None = None,
    ocr_backend: OCRBackend | None = None,
    ocr_dpi: int = 200,
) -> dict[str, Any]:
    if len(content) > MAX_FILE_BYTES:
        raise DocumentParseError(f"file exceeds {MAX_FILE_BYTES} byte parse limit")
    suffix = Path(file_name).suffix.lower()
    kind = suffix.lstrip(".") or (file_type or "").lower()
    if kind == "pdf":
        blocks, summary = _parse_pdf(content, file_name, file_type, ocr_backend, ocr_dpi)
    elif kind == "docx":
        blocks, summary = _parse_docx(content)
    elif kind in {"xlsx", "xlsm"}:
        blocks, summary = _parse_xlsx(content)
    elif kind in {"txt", "md", "csv"}:
        blocks, summary = _parse_text(content)
    else:
        raise DocumentParseError(f"unsupported document type: {kind or 'unknown'}")
    _ensure_limits(content, blocks)
    character_count = sum(len(item["content_text"]) for item in blocks)
    page_assets = (
        render_pdf_page_assets(content, summary["drawing_page_details"])
        if kind == "pdf"
        else []
    )
    return {
        "blocks": blocks,
        "page_assets": page_assets,
        "summary": {
            **summary,
            "file_type": kind,
            "block_count": len(blocks),
            "character_count": character_count,
        },
    }
