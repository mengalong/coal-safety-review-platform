from __future__ import annotations

import os
from io import BytesIO
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFError, TTFont
from reportlab.platypus import KeepTogether, PageBreak, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle


def _summary(content: dict) -> dict:
    task = content.get("task") or {}
    issue_summary = content.get("issue_summary") or {}
    execution = content.get("execution_summary") or {}
    return {
        "报告编号": task.get("task_no", ""),
        "客户名称": task.get("customer_name", ""),
        "产品名称": task.get("product_name", ""),
        "产品型号": task.get("product_model", ""),
        "审核轮次": f"第 {(content.get('round') or {}).get('round_no', 1)} 轮",
        "审核结论": content.get("conclusion", ""),
        "规则执行": f"{execution.get('total', 0)} 条",
        "问题汇总": f"共 {issue_summary.get('total', 0)} 条，已确认 {issue_summary.get('confirmed', 0)} 条",
    }


def _set_cell_shading(cell, fill: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    cell._tc.get_or_add_tcPr().append(shading)


def _set_east_asia_font(run, name: str = "宋体") -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)


def _add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("第 ")
    field_begin = OxmlElement("w:fldChar")
    field_begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = "PAGE"
    field_end = OxmlElement("w:fldChar")
    field_end.set(qn("w:fldCharType"), "end")
    run._r.extend([field_begin, instruction, field_end])
    run.add_text(" 页")
    _set_east_asia_font(run)


def render_docx(content: dict) -> bytes:
    document = Document()
    section = document.sections[0]
    section.start_type = WD_SECTION.NEW_PAGE
    section.page_width, section.page_height = Cm(21), Cm(29.7)
    section.top_margin = section.bottom_margin = Cm(2.2)
    section.left_margin = section.right_margin = Cm(2.4)

    normal = document.styles["Normal"]
    normal.font.name = "宋体"
    normal.font.size = Pt(10.5)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    for style_name in ("Title", "Heading 1", "Heading 2"):
        style = document.styles[style_name]
        style.font.name = "黑体"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    document.styles["Title"].font.size = Pt(20)

    header = section.header.paragraphs[0]
    header.text = "煤矿安标技术文档智能审核平台"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in header.runs:
        _set_east_asia_font(run)
        run.font.size = Pt(9)
    _add_page_number(section.footer.paragraphs[0])

    title = document.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run(str(content.get("title") or "煤矿安标技术文档审核报告"))
    document.add_paragraph()

    summary = _summary(content)
    table = document.add_table(rows=0, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    pairs = list(summary.items())
    for index in range(0, len(pairs), 2):
        row = table.add_row().cells
        for offset, (label, value) in enumerate(pairs[index:index + 2]):
            label_cell, value_cell = row[offset * 2], row[offset * 2 + 1]
            label_cell.text, value_cell.text = label, str(value)
            label_cell.vertical_alignment = value_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            _set_cell_shading(label_cell, "EAF0F5")
            for run in label_cell.paragraphs[0].runs:
                run.bold = True
                _set_east_asia_font(run, "黑体")

    document.add_paragraph("审核问题明细", style="Heading 1")
    issues = content.get("issues") or []
    if not issues:
        document.add_paragraph("本轮审核未形成问题记录。")
    else:
        issue_table = document.add_table(rows=1, cols=5)
        issue_table.style = "Table Grid"
        issue_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        headings = ("序号", "问题编号", "问题标题", "严重度", "人工状态")
        for cell, heading in zip(issue_table.rows[0].cells, headings, strict=True):
            cell.text = heading
            _set_cell_shading(cell, "D9E7EC")
            for run in cell.paragraphs[0].runs:
                run.bold = True
                _set_east_asia_font(run, "黑体")
        for index, issue in enumerate(issues, start=1):
            cells = issue_table.add_row().cells
            values = (index, issue.get("issue_code", ""), issue.get("title", ""), issue.get("severity", ""), issue.get("status", ""))
            for cell, value in zip(cells, values, strict=True):
                cell.text = str(value)
            detail = issue_table.add_row().cells
            detail[0].merge(detail[-1])
            detail[0].text = f"问题说明：{issue.get('description', '')}\n人工结论：{issue.get('manual_conclusion') or '待最终确认'}"

    document.add_paragraph("标准覆盖与执行说明", style="Heading 1")
    standards = content.get("standards") or []
    if standards:
        for item in standards:
            document.add_paragraph(
                f"{item.get('standard_code') or item.get('full_code') or ''} {item.get('standard_name') or ''}",
                style="List Bullet",
            )
    else:
        document.add_paragraph("本报告未包含标准快照。")
    document.add_paragraph("本报告中的模型与规则结论均已进入人工复核流程，最终结论以审核人员确认记录为准。")

    for paragraph in document.paragraphs:
        paragraph.paragraph_format.space_after = Pt(6)
        for run in paragraph.runs:
            _set_east_asia_font(run, "黑体" if paragraph.style.name.startswith(("Title", "Heading")) else "宋体")

    output = BytesIO()
    document.save(output)
    return output.getvalue()


PDF_FONT_NAME = "CoalCJK"


def _register_pdf_font() -> str:
    if PDF_FONT_NAME in pdfmetrics.getRegisteredFontNames():
        return PDF_FONT_NAME
    configured = os.getenv("COAL_REPORT_FONT_PATH")
    candidates = [
        configured,
        "/usr/share/fonts/truetype/wqy/wqy-zenhei.ttc",
        "/usr/share/fonts/truetype/arphic/uming.ttc",
        "/System/Library/Fonts/STHeiti Light.ttc",
        "C:/Windows/Fonts/simsun.ttc",
    ]
    errors = []
    for candidate in filter(None, candidates):
        if not Path(candidate).is_file():
            continue
        try:
            pdfmetrics.registerFont(TTFont(PDF_FONT_NAME, candidate, subfontIndex=0))
            return PDF_FONT_NAME
        except TTFError as exc:
            errors.append(f"{candidate}: {exc}")
    detail = "; ".join(errors) or "no supported font file found"
    raise RuntimeError(f"a TrueType CJK font is required for PDF reports: {detail}")


def _pdf_styles():
    font_name = _register_pdf_font()
    styles = getSampleStyleSheet()
    body = ParagraphStyle("ChineseBody", parent=styles["BodyText"], fontName=font_name, fontSize=10, leading=16, wordWrap="CJK")
    title = ParagraphStyle("ChineseTitle", parent=body, fontSize=20, leading=28, alignment=TA_CENTER, spaceAfter=12)
    heading = ParagraphStyle("ChineseHeading", parent=body, fontSize=14, leading=22, spaceBefore=12, spaceAfter=8, textColor=colors.HexColor("#173B45"))
    small = ParagraphStyle("ChineseSmall", parent=body, fontSize=8.5, leading=13, textColor=colors.HexColor("#53656D"))
    return body, title, heading, small


def render_pdf(content: dict) -> bytes:
    body, title, heading, small = _pdf_styles()
    output = BytesIO()
    document = SimpleDocTemplate(
        output,
        pagesize=A4,
        rightMargin=22 * mm,
        leftMargin=22 * mm,
        topMargin=24 * mm,
        bottomMargin=22 * mm,
        title=str(content.get("title") or "煤矿安标技术文档审核报告"),
        author="煤矿安标技术文档智能审核平台",
    )

    def page(canvas, doc) -> None:
        canvas.saveState()
        canvas.setFont(PDF_FONT_NAME, 8)
        canvas.setFillColor(colors.HexColor("#60727B"))
        canvas.drawString(22 * mm, A4[1] - 14 * mm, "煤矿安标技术文档智能审核平台")
        canvas.drawRightString(A4[0] - 22 * mm, 12 * mm, f"第 {doc.page} 页")
        canvas.restoreState()

    story = [Paragraph(str(content.get("title") or "煤矿安标技术文档审核报告"), title), Spacer(1, 5 * mm)]
    summary_rows = []
    pairs = list(_summary(content).items())
    for index in range(0, len(pairs), 2):
        row = []
        for label, value in pairs[index:index + 2]:
            row.extend((Paragraph(label, small), Paragraph(str(value), body)))
        summary_rows.append(row)
    summary_table = Table(summary_rows, colWidths=[24 * mm, 52 * mm, 24 * mm, 52 * mm], repeatRows=0)
    summary_table.setStyle(TableStyle([
        ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#BFCED3")),
        ("BACKGROUND", (0, 0), (0, -1), colors.HexColor("#EAF0F5")),
        ("BACKGROUND", (2, 0), (2, -1), colors.HexColor("#EAF0F5")),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("LEFTPADDING", (0, 0), (-1, -1), 6),
        ("RIGHTPADDING", (0, 0), (-1, -1), 6),
        ("TOPPADDING", (0, 0), (-1, -1), 7),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 7),
    ]))
    story.extend((summary_table, Paragraph("审核问题明细", heading)))

    issues = content.get("issues") or []
    if not issues:
        story.append(Paragraph("本轮审核未形成问题记录。", body))
    for index, issue in enumerate(issues, start=1):
        issue_header = Table(
            [[Paragraph(f"{index}. {issue.get('title', '')}", body), Paragraph(str(issue.get("severity", "")), small)]],
            colWidths=[132 * mm, 20 * mm],
        )
        issue_header.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#EDF3FF")),
            ("BOX", (0, 0), (-1, -1), 0.6, colors.HexColor("#B9C9E8")),
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ("LEFTPADDING", (0, 0), (-1, -1), 7),
            ("TOPPADDING", (0, 0), (-1, -1), 6),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
        ]))
        details = Paragraph(
            f"问题编号：{issue.get('issue_code', '')}<br/>问题说明：{issue.get('description', '')}<br/>人工结论：{issue.get('manual_conclusion') or '待最终确认'}",
            body,
        )
        story.append(KeepTogether([issue_header, Spacer(1, 2 * mm), details, Spacer(1, 4 * mm)]))

    story.extend((Paragraph("标准覆盖与执行说明", heading), Paragraph("本报告固化生成时的标准、规则、执行器、模型版本和人工结论快照。模型与规则结论均需人工最终确认。", body)))
    if len(story) > 24:
        story.insert(24, PageBreak())
    document.build(story, onFirstPage=page, onLaterPages=page)
    return output.getvalue()
