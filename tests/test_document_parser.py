from io import BytesIO

import pytest
from docx import Document
from openpyxl import Workbook
from PIL import Image
from pypdf import PdfWriter

from coal_platform.document_parser import DocumentParseError, parse_document
from coal_platform.ocr import OCRLine, TesseractOCRBackend
from coal_platform.report_renderer import render_pdf


def _docx_bytes() -> bytes:
    document = Document()
    document.add_heading("Technical specification", level=1)
    document.add_paragraph("Model: DSJ80/40/2x75")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Parameter"
    table.cell(0, 1).text = "Value"
    table.cell(1, 0).text = "Voltage"
    table.cell(1, 1).text = "1140 V"
    output = BytesIO()
    document.save(output)
    return output.getvalue()


def _xlsx_bytes() -> bytes:
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Parameters"
    sheet.append(["Name", "Value", "Unit"])
    sheet.append(["Power", 75, "kW"])
    output = BytesIO()
    workbook.save(output)
    workbook.close()
    return output.getvalue()


class FakeOCRBackend:
    engine_name = "fake-ocr:chi_sim+eng"

    def recognize(self, image_content: bytes) -> list[OCRLine]:
        assert image_content.startswith(b"\x89PNG")
        return [
            OCRLine(
                text="矿用隔爆型开关",
                confidence=0.92,
                left=10,
                top=20,
                width=80,
                height=15,
            )
        ]


def test_parse_docx_preserves_paragraph_and_table_row_references() -> None:
    parsed = parse_document(_docx_bytes(), "specification.docx")

    assert parsed["summary"]["parser"] == "python-docx"
    assert parsed["summary"]["table_row_count"] == 2
    assert any(item["content_text"] == "Model: DSJ80/40/2x75" for item in parsed["blocks"])
    assert any(item["source_ref"] == "docx:table:1:row:2" for item in parsed["blocks"])


def test_parse_xlsx_preserves_sheet_and_cell_range() -> None:
    parsed = parse_document(_xlsx_bytes(), "parameters.xlsx")

    assert parsed["summary"]["sheet_names"] == ["Parameters"]
    assert parsed["blocks"][1]["content_text"] == "Power\t75\tkW"
    assert parsed["blocks"][1]["source_ref"] == "xlsx:Parameters!A2:C2"


def test_parse_pdf_extracts_text_and_marks_blank_pages_for_ocr() -> None:
    text_pdf = render_pdf({"title": "Safety report", "task": {"task_no": "TASK-001"}})
    parsed = parse_document(text_pdf, "report.pdf")
    assert parsed["summary"]["page_count"] == 1
    assert parsed["summary"]["needs_ocr"] is False
    assert "Safety report" in "".join(item["content_text"] for item in parsed["blocks"])

    writer = PdfWriter()
    writer.add_blank_page(width=100, height=100)
    output = BytesIO()
    writer.write(output)
    blank = parse_document(output.getvalue(), "scan.pdf")
    assert blank["summary"]["needs_ocr"] is True
    assert blank["summary"]["empty_text_pages"] == [1]

    drawing = parse_document(text_pdf, "product-drawing.pdf")
    assert drawing["summary"]["drawing_pages"] == [1]
    assert drawing["page_assets"][0]["is_drawing"] is True
    assert drawing["page_assets"][0]["content"].startswith(b"\x89PNG")


def test_parse_pdf_uses_ocr_for_blank_page_and_converts_bbox_to_pdf_points() -> None:
    writer = PdfWriter()
    writer.add_blank_page(width=100, height=200)
    output = BytesIO()
    writer.write(output)

    parsed = parse_document(output.getvalue(), "scan.pdf", ocr_backend=FakeOCRBackend(), ocr_dpi=200)

    assert parsed["summary"]["needs_ocr"] is False
    assert parsed["summary"]["ocr_page_count"] == 1
    assert parsed["summary"]["ocr_block_count"] == 1
    block = parsed["blocks"][0]
    assert block["content_text"] == "矿用隔爆型开关"
    assert block["confidence"] == 0.92
    assert block["source_ref"] == "ocr:fake-ocr:chi_sim+eng:page:1:line:1"
    assert block["bbox"]["unit"] == "pt"
    assert block["bbox"]["page_width"] == 100.0
    assert block["bbox"]["page_height"] == 200.0
    assert 0 < block["bbox"]["x"] < 100
    assert 0 < block["bbox"]["y"] < 200


def test_tesseract_backend_groups_words_and_filters_low_confidence(monkeypatch: pytest.MonkeyPatch) -> None:
    data = {
        "text": ["矿用", "开关", "noise"],
        "conf": ["90", "80", "10"],
        "block_num": [1, 1, 1],
        "par_num": [1, 1, 1],
        "line_num": [1, 1, 2],
        "left": [10, 50, 2],
        "top": [20, 20, 2],
        "width": [30, 40, 5],
        "height": [12, 14, 5],
    }
    monkeypatch.setattr("coal_platform.ocr.pytesseract.image_to_data", lambda *args, **kwargs: data)
    image = Image.new("RGB", (120, 80), "white")
    output = BytesIO()
    image.save(output, format="PNG")

    lines = TesseractOCRBackend(minimum_confidence=0.35).recognize(output.getvalue())

    assert lines == [OCRLine("矿用 开关", 0.85, 10, 20, 80, 14)]


def test_parse_document_rejects_unsupported_type() -> None:
    with pytest.raises(DocumentParseError, match="unsupported document type"):
        parse_document(b"binary", "drawing.dwg")
