from io import BytesIO

import pytest
from docx import Document
from pypdf import PdfReader

from coal_platform.report_renderer import render_docx, render_pdf


def test_report_renderers_generate_real_docx_and_pdf_bytes() -> None:
    content = {
        "title": "煤矿安标技术文档审核报告",
        "task": {"task_no": "SH-2026-000001", "customer_name": "测试企业", "product_model": "KBZ-500"},
        "conclusion": "through",
        "issue_summary": {"total": 1, "confirmed": 0},
        "execution_summary": {"total": 2},
        "issues": [{"title": "资料缺失", "severity": "一般", "description": "缺少试验报告"}],
    }
    docx = render_docx(content)
    document = Document(BytesIO(docx))
    assert document.sections[0].page_width.mm == pytest.approx(210, abs=0.1)
    assert any("煤矿安标技术文档审核报告" in paragraph.text for paragraph in document.paragraphs)
    assert any("资料缺失" in cell.text for table in document.tables for row in table.rows for cell in row.cells)
    pdf = render_pdf(content)
    assert pdf.startswith(b"%PDF-")
    reader = PdfReader(BytesIO(pdf))
    assert len(reader.pages) == 1
    text = reader.pages[0].extract_text()
    assert "煤矿安标技术文档审核报告" in text
    assert "资料缺失" in text
